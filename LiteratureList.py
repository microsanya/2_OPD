#api_key = "AIzaSyAHsGmRT8bOhXNPcA6Gb-A8QjzBXAaSd4Q" для гугл апи

import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
from datetime import datetime
import re

class ReferenceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Добавление источников в Word-документ")
        self.root.geometry("800x400")  # Установка размера окна

        # Тип источника
        self.source_type_label = tk.Label(root, text="Тип источника:")
        self.source_type_label.pack()
        self.source_type_var = tk.StringVar()
        self.source_type_entry = tk.Entry(root,textvariable=self.source_type_var)
        self.source_type_entry.pack()

        # Автор
        self.author_label = tk.Label(root, text="Автор/Номер ГОСТа:")
        self.author_label.pack()
        self.author_var = tk.StringVar()
        self.author_entry = tk.Entry(root, textvariable=self.author_var)
        self.author_entry.pack()

        # Название
        self.title_label = tk.Label(root, text="Название/Полное название ГОСТа:")
        self.title_label.pack()
        self.title_var = tk.StringVar()
        self.title_entry = tk.Entry(root, textvariable=self.title_var)
        self.title_entry.pack()

        tk.Label(self.root, text="URL ресурса:").pack()
        self.url_entry = tk.Entry(self.root, width=50)
        self.url_entry.pack()

        tk.Label(self.root, text="Дата обращения:").pack()
        self.access_date_entry = tk.Entry(self.root, width=15)
        self.access_date_entry.pack()

        # Выбор пути для сохранения файла
        self.save_path_button = tk.Button(root, text="Выбрать путь для сохранения файла", command=self.select_save_path)
        self.save_path_button.pack()

        # Сохранить
        self.save_button = tk.Button(root, text="Сохранить", command=self.save_references)
        self.save_button.pack()

        # Новая метка с информацией для пользователя
        info_text = (
            "Для добавления источника информации напишите один из типов: Книга, Статья в сборнике, "
            "Статья в журнале, Диссертация/автореферат диссертации, ГОСТ, Авторское свидетельство,"
            "Патент, Электронный ресурс локального доступа, Электронный ресурс удалённого доступа."
            "В поле автор нужно вписать 1-го автора, даже если их несколько. Удачи с вашей работой! "
            "Почта для помощи и обращений: your-email@example.com."
        )
        self.info_label = tk.Label(root, text=info_text, wraplength=750, justify='center')
        self.info_label.pack(pady=20)

        self.save_path = None
        self.references = []

    def select_save_path(self):
        self.save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if self.save_path:
            messagebox.showinfo("Путь для сохранения выбран", f"Файл будет сохранен по пути: {self.save_path}")

    def save_references(self):
        if not self.save_path:
            messagebox.showerror("Ошибка", "Сначала выберите путь для сохранения файла")
            return

        doc = Document()

        # Добавление заголовка "Список использованных источников"
        heading = doc.add_paragraph()
        run = heading.add_run("Список использованных источников")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Добавление нумерованных источников
        for i, ref in enumerate(self.references, start=1):
            source = f"{i}. {ref}"
            para = doc.add_paragraph(source)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.paragraph_format.left_indent = Cm(1.25)
            para.paragraph_format.line_spacing = 1.5
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

        # Настройка полей страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(3)

        try:
            doc.save(self.save_path)
            messagebox.showinfo("Готово", "Источники успешно добавлены и документ сохранен")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить документ: {e}")

    def add_reference(self):
        source_type = self.source_type_var.get()
        author = self.author_var.get()
        title = self.title_var.get()
        url = self.url_entry.get()  # Извлечение URL из текстового поля
        access_date = self.access_date_entry.get()  # Извлечение даты обращения из текстового поля
        if not source_type or not author or not title or not url or not access_date:
            messagebox.showerror("Ошибка", "Заполните все поля")
            return

        if source_type.lower() == "книга":
            reference = self.get_book_reference(author, title)
        elif source_type.lower() == "статья в сборнике":
            reference = self.get_collection_article_reference(author, title)
        elif source_type.lower() == "статья в журнале":
            reference = self.get_journal_article_reference(author, title)
        elif source_type.lower() == "диссертация/автореферат диссертации":
            reference = self.get_thesis_reference(author, title)
        elif source_type.lower() == "гост":
            reference = self.get_gost_reference(author, title)
        elif source_type.lower() == "авторское свидетельство":
            reference = self.get_patent_reference(author, title)
        elif source_type.lower() == "патент":
            reference = self.get_patent2_reference(author, title)
        elif source_type.lower() == "электронный ресурс локального доступа":
            reference = self.get_elres_local(author, title)
        elif source_type.lower() == "электронный ресурс удалённого доступа":
            reference = self.get_remote_resource_reference(author, title, url, access_date)
        else:
            reference = f"{source_type} - {author} - {title}"
        self.references.append(reference)
        messagebox.showinfo("Добавлено", "Источник добавлен в список")

        # Очистка полей ввода
        self.source_type_var.set("")
        self.author_var.set("")
        self.title_var.set("")

    # Книга
    def get_book_reference(self, author, title):
        api_key = "AIzaSyAHsGmRT8bOhXNPcA6Gb-A8QjzBXAaSd4Q"
        url = f"https://www.googleapis.com/books/v1/volumes?q=intitle:{title}+inauthor:{author}&key={api_key}"
        response = requests.get(url)
        if response.status_code != 200:
            return f"{author} {title} Город: Издательство, Год Кол-во страниц с.(Не удалось получить информацию)"

        data = response.json()
        if "items" not in data or len(data["items"]) == 0:
            return f"{author} {title} Город: Издательство, Год Кол-во страниц с.(Не удалось получить информацию)"

        book = data["items"][0]["volumeInfo"]
        authors = book.get("authors", [])
        title = book.get("title", "")
        publisher = book.get("publisher", "")
        published_date = book.get("publishedDate", "")
        page_count = book.get("pageCount", "")

        # Извлечение года из даты публикации
        year = published_date.split("-")[0] if published_date else ""

        if len(authors) == 1:
            authors_str = authors[0]
        elif 1 < len(authors) <= 3:
            authors_str = ", ".join(authors)
        elif len(authors) > 3:
            authors_str = f"{authors[0]} [и др.]"
        else:
            authors_str = ""

        # Попытка получить город издания из API Google Books
        city = ""
        if "publishedCity" in book:
            city = book["publishedCity"]
        else:
            # Если город издания не указан, пытаемся найти информацию о книге в других источниках
            city = self.get_city_from_other_sources(author, title)

        # Если город издания не найден, используем город по умолчанию
        if not city:
            city = "Москва"  # Можно выбрать другой город по умолчанию

        reference = f"{authors_str}. {title}. {city}: {publisher}, {year}. {page_count} c."
        return reference

    # Статья в сборнике
    def get_collection_article_reference(self, author, title):
        url = f"https://api.crossref.org/works?query.author={author}&query.title={title}&rows=1"
        response = requests.get(url)
        if response.status_code != 200:
            return f"Статья в сборнике - {author} - {title} (Не удалось получить информацию)"

        data = response.json()
        if "message" not in data or "items" not in data["message"] or len(data["message"]["items"]) == 0:
            return f"{author} {title} // Название сборника: доп. информация, например сб. научных трудов, материалы докл. " \
                   f"и пр., с указанием кол-ва частей, если есть Кол-во частей Город: Издательство, Номер тома/выпуска и тп " \
                   f"С. страницы, на которых расположена статья (Информация не найдена)"

        item = data["message"]["items"][0]
        authors = item.get("author", [])
        title = item.get("title", [""])[0]
        container_title = item.get("container-title", [""])[0]
        publisher = item.get("publisher", "")
        published_date = item.get("published-print", {}).get("date-parts", [[None]])[0][0] or ""
        volume = item.get("volume", "")
        issue = item.get("issue", "")
        page = item.get("page", "")
        place = self.get_city_from_other_sources(author, title)

        if len(authors) == 1:
            authors_str = f"{authors[0]['family']} {authors[0]['given']}"
        elif 1 < len(authors) <= 3:
            authors_str = ", ".join([f"{a['family']} {a['given']}" for a in authors])
        elif len(authors) > 3:
            authors_str = f"{authors[0]['family']} {authors[0]['given']} [и др.]"
        else:
            authors_str = ""

        if not place:
            place = "Москва"  # Можно выбрать другой город по умолчанию

        volume_info = f"Т. {volume}, В. {issue}" if volume and issue else ""
        reference = f"{authors_str} {title} // {container_title}: сб. научных трудов. {volume_info} {place}: {publisher}, " \
                    f"{published_date}. C. {page}."
        return reference

    # Статья в журнале
    def get_journal_article_reference(self, author, title):
        url = f"https://api.crossref.org/works?query.author={author}&query.title={title}&filter=type:journal-article&rows=1"
        response = requests.get(url)
        if response.status_code != 200:
            return f"{author} {title}. // Название журнала. Год. №Части и кол-во частей. С. Страницы статьи. " \
                   f"(Не удалось получить информацию)"

        data = response.json()
        if "message" not in data or "items" not in data["message"] or len(data["message"]["items"]) == 0:
            return f"{author} {title}. // Название журнала. Год. №Части и кол-во частей. С. Страницы статьи. " \
                   f"(Не удалось получить информацию)"

        item = data["message"]["items"][0]
        authors = item.get("author", [])
        title = item.get("title", [""])[0]
        container_title = item.get("container-title", [""])[0]
        published_date = item.get("published-print", {}).get("date-parts", [[None]])[0][0] or ""
        issue = item.get("issue", "")
        page = item.get("page", "")

        if len(authors) == 1:
            authors_str = f"{authors[0]['family']} {authors[0]['given']}"
        elif 1 < len(authors) <= 3:
            authors_str = ", ".join([f"{a['family']} {a['given']}" for a in authors])
        elif len(authors) > 3:
            authors_str = f"{authors[0]['family']} {authors[0]['given']} [и др.]"
        else:
            authors_str = ""

        reference = f"{authors_str} {title} // {container_title}. {published_date}. № {issue}. С. {page}."
        return reference

    # Диссертация/автореферат диссертации
    def get_thesis_reference(self, author, title):
        # Формируем URL для запроса к API dissercat.com
        query = f"{author} {title}".replace(" ", "+")
        url = f"https://www.dissercat.com/api/disser?q={query}"

        # Отправляем GET-запрос
        response = requests.get(url)

        # Проверяем статус ответа
        if response.status_code != 200:
            return f"{author} {title}. доп. информация, диссертация это или автореферат, учёная степень автора " \
                   f"(магистр, доктор, кандидат) " \
                   f": дис. … канд. техн. наук.  или : автореф. на соиск. ученой степ. канд. техн. наук. " \
                   f" Город: Год. Кол-во страниц с.(Информация не найдена)"

        # Получаем JSON-ответ
        data = response.json()

        # Проверяем наличие результатов
        if "items" not in data or len(data["items"]) == 0:
            return f"{author} {title}. доп. информация, диссертация это или автореферат, учёная степень автора " \
                   f"(магистр, доктор, кандидат) " \
               f": дис. … канд. техн. наук.  или : автореф. на соиск. ученой степ. канд. техн. наук. " \
               f" Город: Год. Кол-во страниц с.(Информация не найдена)"

        # Получаем первый результат
        result = data["items"][0]

        # Извлекаем необходимую информацию
        author = result.get("author", "")
        title = result.get("title", "")
        degree = result.get("degree", "")
        city = result.get("city", "")
        year = result.get("year", "")
        pages = result.get("pages", "")

        # Формируем строку ссылки
        reference = f"Диссертация/автореферат диссертации - {author}. {title}: {degree}. {city}, {year}. {pages} с."

        return reference

    # ГОСТ
    def get_gost_reference(gost_number, gost_title):
        # Формируем запрос к API ГОСТинфо
        url = f"https://api.gostinfo.ru/gosts?q={gost_number} {gost_title}"
        response = requests.get(url)
        if response.status_code != 200:
            return f"ГОСТ {gost_number}. {gost_title} (Не удалось получить информацию)"

        data = response.json()
        if "items" not in data or len(data["items"]) == 0:
            return f"ГОСТ {gost_number}. {gost_title} (Информация не найдена)"

        gost_info = data["items"][0]  # Предполагается, что API возвращает информацию о первом найденном ГОСТе
        city = gost_info.get("city", "")  # Получаем информацию о городе издания
        publisher = gost_info.get("publisher", "")  # Получаем информацию об издательстве
        year = gost_info.get("year", "")  # Получаем информацию о годе издания
        pages = gost_info.get("pages", "")  # Получаем информацию о количестве страниц

        # Оформляем информацию о ГОСТе в виде шаблона
        reference = f"{gost_number}. {gost_title}. "
        if city.lower() == "москва":
            reference += "М.: "
        elif city.lower() == "санкт-петербург":
            reference += "СПб.: "
        else:
            reference += f"{city}: "
        reference += f"{publisher}, {year}. {pages} с."

        return reference

    # Авторское свидетельство
    def get_patent_reference(self, author, title):
        api_key = "AIzaSyAHsGmRT8bOhXNPcA6Gb-A8QjzBXAaSd4Q"
        search_engine_id = "b3c479e79bbda4b33"

        query = f"{title} {author}"
        url = f"https://www.googleapis.com/customsearch/v1?q={query}&cx={search_engine_id}&key={api_key}"

        response = requests.get(url)
        if response.status_code != 200:
            return f"{title} а. с. Номер авторского свидетельства Страна №  Код страны Номер заявки А1" \
                   f"{author}; заявл. Дата заявления; опубл. Дата публикации (Не удалось получить информацию)"

        data = response.json()
        if "items" not in data or len(data["items"]) == 0:
            return f"{title} а. с. Номер авторского свидетельства Страна №  Код страны Номер заявки А1" \
                   f"{author}; заявл. Дата заявления; опубл. Дата публикации (Не удалось получить информацию)"

        # Обработка данных о патенте
        item = data["items"][0]
        snippet = item["snippet"]

        # Регулярные выражения для извлечения информации
        title_pattern = re.compile(r"Название:\s*(.*)")
        patent_number_pattern = re.compile(r"а\.\s*с\.\s*(\d+)")
        application_number_pattern = re.compile(r"№\s*SU\s*(\d+)\s*A1")
        authors_pattern = re.compile(r"Авторы:\s*(.*)")
        filing_date_pattern = re.compile(r"заявл\.\s*(\d{2}\.\d{2}\.\d{2})")
        publication_date_pattern = re.compile(r"опубл\.\s*(\d{2}\.\d{2}\.\d{2})")

        # Поиск соответствий в snippet
        title_match = title_pattern.search(snippet)
        patent_number_match = patent_number_pattern.search(snippet)
        application_number_match = application_number_pattern.search(snippet)
        authors_match = authors_pattern.search(snippet)
        filing_date_match = filing_date_pattern.search(snippet)
        publication_date_match = publication_date_pattern.search(snippet)

        title = title_match.group(1) if title_match else title
        patent_number = patent_number_match.group(1) if patent_number_match else ""
        application_number = application_number_match.group(1) if application_number_match else ""
        authors = authors_match.group(1).split(", ") if authors_match else [author]
        filing_date = filing_date_match.group(1) if filing_date_match else ""
        publication_date = publication_date_match.group(1) if publication_date_match else ""

        # Форматирование строки ссылки
        authors_str = ", ".join([f"{a}" for a in authors])
        if len(authors) > 3:
            authors_str = f"{authors[0]} [и др.]"

        reference = (f"{title}: а. с. {patent_number} СССР. "
                     f"№ SU {application_number} A1 "
                     f"/ {authors_str}; заявл. {filing_date}; опубл. {publication_date}.")
        return reference

    # Патент
    def get_patent2_reference(author, title):
        base_url = "https://patentsview.org/patents/query"

        # Подготовка запроса к API
        query = {"q": {"patent_title": title, "inventor_last_name": author.split()[0]},
                 "f": ["patent_number", "patent_title", "patent_date", "inventor_first_name", "inventor_last_name",
                       "application_number", "filing_date", "publication_date"]}

        response = requests.get(base_url, params={"q": query})

        if response.status_code == 200:
            data = response.json()
            if data['patents']:
                patent = data['patents'][0]
                patent_title = patent.get("patent_title", title)
                patent_number = patent.get("patent_number", "")
                application_number = patent.get("application_number", "")
                filing_date = patent.get("filing_date", "")
                publication_date = patent.get("publication_date", "")
                inventors = [f"{inv['inventor_first_name'][0]}. {inv['inventor_last_name']}" for inv in
                             patent.get("inventors", [])]
                inventors_str = ", ".join(inventors)
                if len(inventors) > 3:
                    inventors_str = f"{inventors[0]} [и др.]"

                # Форматирование по шаблону
                reference = f"{patent_title}: пат. {patent_number} Рос. Федерация. № {application_number} / {inventors_str}.;" \
                            f" заявл. {filing_date}; опубл. {publication_date}."
                return reference
            else:
                return f"{title}: пат. Номер патента Рос. Федерация. № Номер заявки {author}; заявл. Дата заявления; " \
                       f" опубл. Дата публикацц(Информация не найдена)"
        else:
            return f"{title}: пат. Номер патента Рос. Федерация. № Номер заявки {author}; заявл. Дата заявления; " \
                       f" опубл. Дата публикацц(Информация не найдена)"

    # Электронный ресурс локального доступа
    def get_elres_local (author, title):
        #анрил найти базу для этого, пользователь будет писать в ворде в шаблон
        reference = f"{author} (если авторов 1 <= x <= 3, иначе пропустить) {title} [Электронный ресурс]" \
                    f"если авторов 4 и более, записывается первый из них. Если авторов нет, записывается редактор / И.О. Фамилия [и др.].  или / под ред. " \
                    f"И.О. Фамилия. Город: Издатель, Год. полное название кол-во носителей n носитель " \
                    f"Сокращённое название носителя"
        return reference

    # Электронный ресурс удалённого доступа
    def get_remote_resource_reference(self, author, title, url, access_date):
        reference = f"{author}. {title}. [Электронный ресурс]: {url} (дата обращения: {access_date})."
        return reference

    # Не всегда в апишках есть город, на всякий еще одну хуйнула
    def get_city_from_other_sources(self, author, title):
        url = f"https://openlibrary.org/search.json?author={author}&title={title}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            if "docs" in data and len(data["docs"]) > 0:
                doc = data["docs"][0]
                publish_places = doc.get("publish_places", [])
                if publish_places:
                    return publish_places[0]  # Возвращаем первое место издания, если оно доступно
        return None

if __name__ == "__main__":
    root = tk.Tk()
    app = ReferenceApp(root)

    # Кнопка для добавления источника в список
    add_button = tk.Button(root, text="Добавить источник", command=app.add_reference)
    add_button.pack()

    root.mainloop()